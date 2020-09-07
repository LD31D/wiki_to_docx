from docx import Document
from random import choice
from wikipedia import page, exceptions, set_lang


def write_to_docx(wiki_page):
	document = Document()

	document.add_heading(wiki_page.title, 0)
	paragraph = document.add_paragraph()

	content = wiki_page.content.split('\n\n')

	for element in content[:-2]:
		paragraph.add_run('\n\n')
		paragraph.add_run(element)

	document.save(f'{wiki_page.title}.docx')


def get_page(query: str):
	set_lang("ru")

	try:
		wiki_page = page(query)

	except exceptions.DisambiguationError as e:
		wiki_page = page(choice(e.options))

	print(wiki_page.url)

	return wiki_page


def main():
	query = input('Input query>> ')

	wiki_page = get_page(query)

	write_to_docx(wiki_page)


if __name__ == '__main__':
	main()
